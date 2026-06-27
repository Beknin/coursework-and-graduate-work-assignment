from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Dict


class WordGenerator:
    """Генератор Word-документов (приказов)"""

    @staticmethod
    def generate_order(data: Dict, output_path: str) -> str:
        """Генерирует приказ в Word"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading(f"{data['order_type']} ПРИКАЗ", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата
        date_paragraph = doc.add_paragraph(f"Дата: {data['date']}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Текст приказа
        doc.add_paragraph("На основании решения кафедры об утверждении тем выпускных квалификационных работ и курсовых проектов, ПРИКАЗЫВАЮ:")
        doc.add_paragraph()
        
        # Таблица со студентами
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        headers = ["№", "Студент", "Тема", "Руководитель"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Данные
        for idx, student in enumerate(data['students'], 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = student.get('student_name', '—')
            row.cells[2].text = student.get('topic_title', '—')
            row.cells[3].text = student.get('teacher_name', '—')
        
        doc.add_paragraph()
        
        # Подписи
        doc.add_paragraph("Заведующий кафедрой: _________________")
        doc.add_paragraph("Секретарь: _________________")
        doc.add_paragraph()
        doc.add_paragraph("С приказом ознакомлены:")
        doc.add_paragraph()
        doc.add_paragraph("_________________")
        
        # Сохраняем
        if not Path(output_path).parent.exists():
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(output_path)
        return output_path