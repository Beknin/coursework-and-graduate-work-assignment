from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from typing import List, Dict
import os


class OrderService:
    """Сервис для генерации приказов в Word"""

    @staticmethod
    def generate_order(data: List[Dict], order_type: str, output_path: str = None) -> str:
        if output_path is None:
            output_path = f"{order_type}_order_{date.today().strftime('%Y%m%d')}.docx"

        doc = Document()

        # Заголовок
        title = doc.add_heading(f"{order_type.upper()} ПРИКАЗ", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата
        p = doc.add_paragraph(f"Дата: {date.today().strftime('%d.%m.%Y')}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()
        doc.add_paragraph("На основании решения кафедры об утверждении тем выпускных квалификационных работ и курсовых проектов, ПРИКАЗЫВАЮ:")
        doc.add_paragraph()

        # Таблица
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        headers = ["№", "Студент", "Группа", "Тема", "Руководитель"]
        for i, hdr in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True

        for idx, row in enumerate(data, 1):
            cells = table.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = row.get('student_name', '—')
            cells[2].text = row.get('group', '—')
            cells[3].text = row.get('topic_title', '—')
            cells[4].text = row.get('teacher_name', '—')

        doc.add_paragraph()
        doc.add_paragraph("Заведующий кафедрой: _________________")
        doc.add_paragraph("Секретарь: _________________")
        doc.add_paragraph("С приказом ознакомлены:")

        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        return output_path